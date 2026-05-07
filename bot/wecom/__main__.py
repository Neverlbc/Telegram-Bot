"""企业微信智能机器人入口：python -m bot.wecom"""

from __future__ import annotations

import asyncio
import logging
import re
import signal
from typing import Any

from bot.config import settings
from bot.logging_config import setup_logging
from bot.wecom.client import WecomBotClient
from bot.wecom.llm import chat_with_tools

logger = logging.getLogger(__name__)


def _extract_user_text(frame: dict[str, Any]) -> str:
    """从 aibot_msg_callback 帧里取出文本内容，去除 @机器人 前缀."""
    body = frame.get("body") or {}
    msgtype = body.get("msgtype")
    if msgtype != "text":
        return ""
    text = (body.get("text") or {}).get("content", "")
    # 群聊 @机器人 时 content 形如 "@机器人名 你好"，去掉 @xxx 前缀
    return re.sub(r"^@\S+\s*", "", text).strip()


async def main_async() -> None:
    setup_logging(settings.log_level, settings.log_format)

    if not settings.wecom_bot_id or not settings.wecom_bot_secret:
        logger.error("缺少 WECOM_BOT_ID / WECOM_BOT_SECRET，wecom-agent 无法启动")
        return

    client: WecomBotClient | None = None

    async def _handle_message_async(
        frame: dict[str, Any], text: str, stream_id: str
    ) -> None:
        """LLM 处理 + 回复，独立 task 跑，不阻塞 WebSocket 主循环."""
        logger.info("[wecom] LLM 开始处理: %r", text[:80])
        try:
            reply = await chat_with_tools(text)
            logger.info("[wecom] LLM 完成 len=%d preview=%r", len(reply), reply[:80])
        except Exception as exc:
            logger.exception("LLM dispatch failed")
            reply = f"❌ 处理消息时出错：{exc}"

        if client is not None:
            try:
                # 同一 stream_id + finish=True，企业微信客户端用最终内容替换占位文字
                await client.reply_text(frame, reply, stream_id=stream_id)
            except Exception:
                logger.exception("回复消息失败")

    async def on_message(frame: dict[str, Any]) -> None:
        body = frame.get("body") or {}
        chat_type = body.get("chattype", "single")
        sender = (body.get("from") or {}).get("userid", "")
        
        # 特殊处理：如果用户发的是 txt 文件（长文本超限被企微转为文件）
        msgtype = body.get("msgtype")
        text = ""
        if msgtype == "text":
            text = _extract_user_text(frame)
        elif msgtype == "attachment" or msgtype == "file":
            file_info = body.get(msgtype, {}) or body.get("file", {}) or body.get("attachment", {})
            file_id = file_info.get("fileid") or file_info.get("media_id") or ""
            file_name = file_info.get("filename") or file_info.get("name") or ""
            # 兼容企微将 docx 作为 file 或 docmsg 等类型发送
            if file_id and (file_name.endswith(".txt") or "=" in file_name or file_name.endswith(".docx") or file_name.endswith(".doc")):
                if client is not None:
                    try:
                        content_bytes = await client.download_media(file_id)
                        if file_name.endswith(".docx"):
                            import io
                            import docx
                            doc_obj = docx.Document(io.BytesIO(content_bytes))
                            text = "\n".join([p.text for p in doc_obj.paragraphs]).strip()
                            logger.info("[wecom] 已将docx附件转换为文本，长度=%d", len(text))
                        elif file_name.endswith(".doc"):
                            await client.reply_text(frame, "⚠️ 抱歉，我不支持老版本的 .doc 格式文件。\n请将文件**另存为 .docx** 或 **.txt** 格式后再发给我！")
                            return
                        else:
                            text = content_bytes.decode('utf-8', errors='ignore').strip()
                            logger.info("[wecom] 已将txt附件转换为文本，长度=%d", len(text))
                    except Exception as e:
                        logger.error("下载/解析附件失败: %s", e)
        elif msgtype in ["doc", "docmsg", "link", "markdown"]:
            file_info = body.get(msgtype, {})
            # 企微在线文档也是 docmsg，但某些情况它也是带 fileid 的真附件
            file_name = file_info.get("title") or file_info.get("doc_title") or ""
            if file_name.endswith(".docx") and file_info.get("fileid"):
                if client is not None:
                    try:
                        content_bytes = await client.download_media(file_info.get("fileid"))
                        import io
                        import docx
                        doc_obj = docx.Document(io.BytesIO(content_bytes))
                        text = "\n".join([p.text for p in doc_obj.paragraphs]).strip()
                        logger.info("[wecom] 已从docmsg中将docx附件转换为文本，长度=%d", len(text))
                    except Exception as e:
                        logger.error("下载/解析docmsg附件失败: %s", e)
            else:
                if client is not None:
                    await client.reply_text(frame, "⚠️ 检测到你发送了在线文档或链接。\n\n由于企微权限限制，机器人**无法直接读取在线文档**的内容。\n👉 **更新Cookie的正确姿势**：\n在电脑桌面新建一个正常的「记事本(TXT)文件」或直接发「Word(.docx)文档」，把内容粘贴进去保存，然后发进聊天框！")
                return
        else:
            return  # 忽略语音、图片等不可读消息

        logger.info("[wecom] msg from=%s chat=%s msgtype=%s text_len=%d", sender, chat_type, msgtype, len(text))
        if not text:
            return

        if client is None:
            return

        # --- 第一层直接指令拦截：手动更新 Cookie ---
        if text.startswith("/update_cookie"):
            parts = text.split(maxsplit=2)
            store_name = parts[1] if len(parts) > 1 else ""
            cookie_str = parts[2] if len(parts) > 2 else ""
            
            if not store_name or not cookie_str:
                await client.reply_text(frame, "❌ 格式错误。正确格式(可在消息内直接发或分开发):\n`/update_cookie [店铺名] [你的超长Cookie]`\n\n> 💡贴士：如果在电脑复制Cookie太长变成文件发送，你可以在txt或Word的第一行写上 `/update_cookie 主店`，换行后再粘贴cookie发送！")
                return
            
            from bot.services.aliexpress_mtop import save_cookie
            save_cookie(store_name, cookie_str)
            await client.reply_text(frame, f"✅ 已成功更新并保存【{store_name}】的授权 Cookie！快试试让 AI 帮你创建折扣码吧。")
            return

        # 立即发占位帧（finish=False），让用户看到"处理中"而非空白等待
        stream_id = client.new_stream_id()
        try:
            await client.reply_stream(frame, stream_id, "🤔 处理中，请稍候…", False)
        except Exception:
            logger.exception("发送占位帧失败")

        # LLM + 最终回复放后台，不阻塞 WS 主循环
        asyncio.create_task(_handle_message_async(frame, text, stream_id))

    async def on_event(frame: dict[str, Any]) -> None:
        body = frame.get("body") or {}
        event_type = body.get("event_type") or body.get("eventtype") or ""
        logger.info("[wecom] event=%s", event_type)
        # 进入会话事件 → 发欢迎语
        if client is not None and event_type in {"enter_chat", "enter", "subscribe"}:
            try:
                await client.reply_welcome(
                    frame,
                    f"您好，我是 {settings.wecom_bot_name}。可以问我：\n"
                    "  • 莫斯科现货库存\n"
                    "  • 今天的机器人日报",
                )
            except Exception:
                logger.exception("发送欢迎语失败")

    client = WecomBotClient(
        bot_id=settings.wecom_bot_id,
        secret=settings.wecom_bot_secret,
        ws_url=settings.wecom_ws_url,
        on_message=on_message,
        on_event=on_event,
    )

    loop = asyncio.get_running_loop()
    stop_event = asyncio.Event()

    def _signal_handler() -> None:
        logger.info("收到退出信号，准备关闭…")
        stop_event.set()

    for sig in (signal.SIGINT, signal.SIGTERM):
        try:
            loop.add_signal_handler(sig, _signal_handler)
        except NotImplementedError:
            pass  # Windows 下可能不支持

    runner = asyncio.create_task(client.run_forever())
    waiter = asyncio.create_task(stop_event.wait())

    done, pending = await asyncio.wait({runner, waiter}, return_when=asyncio.FIRST_COMPLETED)
    for task in pending:
        task.cancel()
    await client.close()


def main() -> None:
    asyncio.run(main_async())


if __name__ == "__main__":
    main()
